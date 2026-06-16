import os
import io
import uuid
import tempfile
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import fitz
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from clause_checker import analyze_risks, parse_chapters, calculate_score, chapter_to_dict

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = tempfile.mkdtemp()
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_path):
    text = ""
    doc = fitz.open(file_path)
    for page in doc:
        text += page.get_text()
    doc.close()
    return text


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text


@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': '未找到文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400

    if not file or not allowed_file(file.filename):
        return jsonify({'error': '不支持的文件格式'}), 400

    filename = file.filename
    file_ext = filename.rsplit('.', 1)[1].lower()
    document_id = str(uuid.uuid4())
    saved_filename = f"{document_id}.{file_ext}"
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], saved_filename)
    file.save(file_path)

    try:
        if file_ext == 'pdf':
            content = extract_text_from_pdf(file_path)
        elif file_ext in ['docx', 'doc']:
            content = extract_text_from_docx(file_path)
        else:
            return jsonify({'error': '不支持的文件格式'}), 400

        title = os.path.splitext(filename)[0]

        os.remove(file_path)

        return jsonify({
            'documentId': document_id,
            'content': content,
            'title': title
        })

    except Exception as e:
        if os.path.exists(file_path):
            os.remove(file_path)
        return jsonify({'error': f'文件解析失败: {str(e)}'}), 500


@app.route('/api/analyze', methods=['POST'])
def analyze_document():
    data = request.get_json()
    content = data.get('content', '')
    title = data.get('title', '未命名文档')

    if not content:
        return jsonify({'error': '文档内容为空'}), 400

    try:
        risks = analyze_risks(content)
        chapters = parse_chapters(content)
        chapters_dict = [chapter_to_dict(ch) for ch in chapters]

        return jsonify({
            'risks': risks,
            'chapters': chapters_dict
        })

    except Exception as e:
        return jsonify({'error': f'分析失败: {str(e)}'}), 500


@app.route('/api/score', methods=['POST'])
def get_score():
    data = request.get_json()
    content = data.get('content', '')
    risks = data.get('risks', [])

    try:
        score = calculate_score(content, risks)
        return jsonify(score)

    except Exception as e:
        return jsonify({'error': f'评分计算失败: {str(e)}'}), 500


@app.route('/api/export', methods=['POST'])
def export_report():
    data = request.get_json()
    original = data.get('original', '')
    modified = data.get('modified', '')
    risks = data.get('risks', [])

    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                leftMargin=2 * cm, rightMargin=2 * cm,
                                topMargin=2 * cm, bottomMargin=2 * cm)

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#165DFF'),
            spaceAfter=20,
            alignment=1
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#1D2129'),
            spaceBefore=15,
            spaceAfter=10,
            borderPadding=(0, 0, 5, 0)
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['BodyText'],
            fontSize=10,
            textColor=colors.HexColor('#4E5969'),
            leading=18,
            spaceAfter=8
        )
        risk_high = ParagraphStyle(
            'RiskHigh',
            parent=normal_style,
            textColor=colors.HexColor('#F53F3F'),
            backColor=colors.HexColor('#FFEDED')
        )

        story = []

        story.append(Paragraph('合同审查报告', title_style))
        story.append(Spacer(1, 20))

        total_score = 100 - sum(8 if r['severity'] == 'high' else 4 if r['severity'] == 'medium' else 1 for r in risks)
        total_score = max(20, total_score)

        score_data = [
            ['综合评分', f'{total_score}/100'],
            ['高风险项', f'{sum(1 for r in risks if r["severity"] == "high")} 项'],
            ['中风险项', f'{sum(1 for r in risks if r["severity"] == "medium")} 项'],
            ['低风险项', f'{sum(1 for r in risks if r["severity"] == "low")} 项'],
        ]

        score_table = Table(score_data, colWidths=[6 * cm, 6 * cm])
        score_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#165DFF')),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#F0F5FF')),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F7F8FA')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E5E6EB')),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('LEFTPADDING', (0, 0), (-1, -1), 15),
            ('RIGHTPADDING', (0, 0), (-1, -1), 15),
        ]))
        story.append(score_table)
        story.append(Spacer(1, 20))

        story.append(Paragraph('一、风险条款详情', heading_style))

        if risks:
            for idx, risk in enumerate(risks, 1):
                severity_label = {'high': '高风险', 'medium': '中风险', 'low': '低风险'}[risk['severity']]
                type_labels = {'penalty': '违约金条款', 'termination': '解约权条款', 'disclaimer': '免责条款', 'other': '其他风险'}
                type_label = type_labels.get(risk['type'], '其他风险')

                story.append(Paragraph(f'{idx}. 【{severity_label}】{type_label}', heading_style))
                story.append(Paragraph(f'<b>风险原文：</b>{risk["text"]}', normal_style))
                story.append(Paragraph(f'<b>风险描述：</b>{risk["description"]}', normal_style))
                story.append(Paragraph(f'<b>修改建议：</b>{risk["suggestion"]}', normal_style))
                story.append(Paragraph(f'<b>法律依据：</b>{risk["legal_basis"]}', normal_style))
                story.append(Spacer(1, 8))
        else:
            story.append(Paragraph('未检测到明显风险条款。', normal_style))

        if original != modified:
            story.append(Spacer(1, 15))
            story.append(Paragraph('二、修改对比', heading_style))

            change_count = abs(len(modified) - len(original))
            if len(modified) > len(original):
                story.append(Paragraph(f'本次修改共新增约 {change_count} 字符。', normal_style))
            else:
                story.append(Paragraph(f'本次修改共删除约 {change_count} 字符。', normal_style))

        story.append(Spacer(1, 20))
        story.append(Paragraph('— 本报告由法律文书智能审查助手自动生成 —', ParagraphStyle(
            'Footer',
            parent=normal_style,
            textColor=colors.HexColor('#86909C'),
            alignment=1,
            fontSize=9
        )))

        doc.build(story)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name='合同审查报告.pdf',
            mimetype='application/pdf'
        )

    except Exception as e:
        return jsonify({'error': f'导出失败: {str(e)}'}), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'ok', 'message': '法律文书智能审查助手服务运行正常'})


if __name__ == '__main__':
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    print('法律文书智能审查助手服务启动中...')
    print('服务地址: http://localhost:5000')
    app.run(host='0.0.0.0', port=5000, debug=True)
